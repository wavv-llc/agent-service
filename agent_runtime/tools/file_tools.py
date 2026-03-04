"""
File manipulation tools — Excel, Word, PDF.

Each tool exposes:
  - A Python function used by executor.py
  - A TOOL_SCHEMA dict for registering with the Anthropic messages API
"""
from __future__ import annotations

from typing import Any

# ---------------------------------------------------------------------------
# Excel (openpyxl)
# ---------------------------------------------------------------------------


def read_excel(path: str, sheet: str | None = None) -> dict:
    """
    Read an Excel file and return its contents as a list of row dicts.
    Requires: openpyxl
    """
    import openpyxl  # lazy import — not available in all containers

    wb = openpyxl.load_workbook(path, read_only=True, data_only=True)
    ws = wb[sheet] if sheet else wb.active

    headers: list[str] = []
    rows: list[dict] = []

    for i, row in enumerate(ws.iter_rows(values_only=True)):
        if i == 0:
            headers = [
                str(c) if c is not None else f'col_{j}' for j, c in enumerate(row)
            ]
        else:
            rows.append(dict(zip(headers, row, strict=False)))

    wb.close()
    return {'sheet': ws.title, 'headers': headers, 'rows': rows, 'row_count': len(rows)}


def write_excel(path: str, data: list[dict]) -> dict:
    """Write *data* (list of row dicts) to a new Excel file at *path*."""
    import openpyxl

    wb = openpyxl.Workbook()
    ws = wb.active

    if not data:
        wb.save(path)
        wb.close()
        return {'path': path, 'rows_written': 0}

    headers = list(data[0].keys())
    ws.append(headers)
    for row in data:
        ws.append([row.get(h) for h in headers])

    wb.save(path)
    wb.close()
    return {'path': path, 'rows_written': len(data)}


EXCEL_READ_SCHEMA: dict = {
    'name': 'file_excel_read',
    'description': 'Read an Excel (.xlsx) file and return its rows as JSON.',
    'input_schema': {
        'type': 'object',
        'required': ['path'],
        'properties': {
            'path': {
                'type': 'string',
                'description': 'Absolute path to the .xlsx file.',
            },
            'sheet': {
                'type': 'string',
                'description': 'Sheet name (optional; defaults to active sheet).',
            },
        },
    },
}

EXCEL_WRITE_SCHEMA: dict = {
    'name': 'file_excel_write',
    'description': 'Write rows to a new Excel (.xlsx) file.',
    'input_schema': {
        'type': 'object',
        'required': ['path', 'data'],
        'properties': {
            'path': {'type': 'string'},
            'data': {'type': 'array', 'items': {'type': 'object'}},
        },
    },
}


# ---------------------------------------------------------------------------
# Word (python-docx)
# ---------------------------------------------------------------------------


def read_word(path: str) -> dict:
    """Extract text from a .docx file."""
    from docx import Document  # python-docx

    doc = Document(path)
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    return {'path': path, 'paragraphs': paragraphs, 'paragraph_count': len(paragraphs)}


def write_word(path: str, paragraphs: list[str]) -> dict:
    """Write *paragraphs* to a new .docx file at *path*."""
    from docx import Document

    doc = Document()
    for text in paragraphs:
        doc.add_paragraph(text)
    doc.save(path)
    return {'path': path, 'paragraphs_written': len(paragraphs)}


WORD_READ_SCHEMA: dict = {
    'name': 'file_word_read',
    'description': 'Read a Word (.docx) file and return its paragraph text.',
    'input_schema': {
        'type': 'object',
        'required': ['path'],
        'properties': {
            'path': {'type': 'string'},
        },
    },
}

WORD_WRITE_SCHEMA: dict = {
    'name': 'file_word_write',
    'description': 'Write paragraphs to a new Word (.docx) file.',
    'input_schema': {
        'type': 'object',
        'required': ['path', 'paragraphs'],
        'properties': {
            'path': {'type': 'string'},
            'paragraphs': {'type': 'array', 'items': {'type': 'string'}},
        },
    },
}


# ---------------------------------------------------------------------------
# PDF (pypdf)
# ---------------------------------------------------------------------------


def read_pdf(path: str) -> dict:
    """Extract text and metadata from a PDF file."""
    from pypdf import PdfReader  # pypdf

    reader = PdfReader(path)
    pages = [page.extract_text() or '' for page in reader.pages]
    meta = {k: str(v) for k, v in (reader.metadata or {}).items()}
    return {
        'path': path,
        'page_count': len(pages),
        'text': '\n\n'.join(pages),
        'metadata': meta,
    }


PDF_READ_SCHEMA: dict = {
    'name': 'file_pdf_read',
    'description': 'Extract text and metadata from a PDF file.',
    'input_schema': {
        'type': 'object',
        'required': ['path'],
        'properties': {
            'path': {'type': 'string'},
        },
    },
}


# ---------------------------------------------------------------------------
# Tool registry (used by executor to dispatch tool calls)
# ---------------------------------------------------------------------------

TOOL_FUNCTIONS: dict[str, Any] = {
    'file_excel_read': read_excel,
    'file_excel_write': write_excel,
    'file_word_read': read_word,
    'file_word_write': write_word,
    'file_pdf_read': read_pdf,
}

ALL_SCHEMAS: list[dict] = [
    EXCEL_READ_SCHEMA,
    EXCEL_WRITE_SCHEMA,
    WORD_READ_SCHEMA,
    WORD_WRITE_SCHEMA,
    PDF_READ_SCHEMA,
]
